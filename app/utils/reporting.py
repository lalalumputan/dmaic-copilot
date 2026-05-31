from __future__ import annotations
import re
import json
import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ══════════════════════════════════════════════════════════════════════════════
# 1.  UTILITY HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _get(d: Dict[str, Any], path: List[str], default=None):
    cur: Any = d
    for k in path:
        if not isinstance(cur, dict) or k not in cur:
            return default
        cur = cur[k]
    return cur


def _as_text(val: Any) -> str:
    if val is None:
        return "(not provided)"
    if isinstance(val, (list, tuple, set)):
        return ", ".join(str(v) for v in val)
    if isinstance(val, dict):
        return "; ".join(f"{k}: {v}" for k, v in val.items())
    return str(val)


def _normalize_define(define_state: Dict[str, Any]) -> Dict[str, Any]:
    inputs  = define_state.get("inputs")  or {}
    outputs = define_state.get("outputs") or {}

    project_name = (
        inputs.get("project_name")
        or define_state.get("project_name")
        or outputs.get("project_name")
        or "Unnamed Project"
    )
    problem       = outputs.get("problem_statement") or define_state.get("problem_statement")
    goal          = outputs.get("goal_statement")    or define_state.get("goal_statement")
    business_case = outputs.get("business_case")     or define_state.get("business_case")

    scope     = outputs.get("project_scope") or define_state.get("project_scope") or {}
    scope_in  = scope.get("in_scope")      or scope.get("inScope")
    scope_out = scope.get("out_of_scope")  or scope.get("outOfScope")

    ctqs      = outputs.get("ctq_list")          or define_state.get("ctq_list")          or []
    sipoc     = outputs.get("sipoc")              or define_state.get("sipoc")              or {}
    early     = outputs.get("early_measure_plan") or define_state.get("early_measure_plan") or {}
    y_var     = outputs.get("y_variable")         or define_state.get("y_variable")
    x_cat     = outputs.get("x_categories")       or define_state.get("x_categories")
    risk_level   = outputs.get("risk_level")      or define_state.get("risk_level")
    project_type = outputs.get("project_type")    or define_state.get("project_type")

    charter_upload = inputs.get("charter_data_uploaded") or {}
    charter_llm    = outputs.get("project_charter") or {}
    charter        = charter_upload if charter_upload.get("available") else charter_llm

    benefit_upload = inputs.get("benefit_data_uploaded") or {}
    benefit_llm    = outputs.get("benefit_estimate") or {}
    benefit        = benefit_upload if benefit_upload.get("available") else benefit_llm
    benefit_from_upload = benefit_upload.get("available", False)

    return {
        "project_name":       project_name,
        "problem":            problem,
        "goal":               goal,
        "business_case":      business_case,
        "scope_in":           scope_in,
        "scope_out":          scope_out,
        "ctqs":               ctqs,
        "sipoc":              sipoc,
        "early_measure_plan": early,
        "y_variable":         y_var,
        "x_categories":       x_cat,
        "risk_level":         risk_level,
        "project_type":       project_type,
        "charter":            charter,
        "benefit":            benefit,
        "benefit_from_upload": benefit_from_upload,
    }


def _bullet_list(val: Any) -> str:
    if val in (None, "", [], {}):
        return "_not available_"
    if isinstance(val, list):
        return "\n".join([f"- {v}" for v in val]) if val else "_not available_"
    if isinstance(val, dict):
        lines = []
        for k, v in val.items():
            lines.append(f"- **{k}**: {_as_text(v)}")
        return "\n".join(lines) if lines else "_not available_"
    return f"- {val}"


def _sipoc_md(sipoc: Any) -> str:
    if not isinstance(sipoc, dict) or not sipoc:
        return "_not available_"
    suppliers = sipoc.get("Suppliers") or sipoc.get("suppliers") or []
    inputs    = sipoc.get("Inputs")    or sipoc.get("inputs")    or []
    process   = sipoc.get("Process")   or sipoc.get("process")   or []
    outputs   = sipoc.get("Outputs")   or sipoc.get("outputs")   or []
    customers = sipoc.get("Customers") or sipoc.get("customers") or []
    return (
        "**Suppliers**\n"  + _bullet_list(suppliers) + "\n\n"
        "**Inputs**\n"     + _bullet_list(inputs)    + "\n\n"
        "**Process**\n"    + _bullet_list(process)   + "\n\n"
        "**Outputs**\n"    + _bullet_list(outputs)   + "\n\n"
        "**Customers**\n"  + _bullet_list(customers)
    )


def _strip_md_heading(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"^\s*#{1,6}\s*Agent Insight\s*\n?", "", text.strip(), flags=re.IGNORECASE)


# ══════════════════════════════════════════════════════════════════════════════
# 2.  DOCUMENT-LEVEL HELPERS  (take `doc` as first arg)
# ══════════════════════════════════════════════════════════════════════════════

def _doc_setup(doc: Document):
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)


def _doc_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor(0x07, 0x59, 0x85)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "0759a0")
    pBdr.append(bot)
    pPr.append(pBdr)


class _SecNum:
    """Dynamic section numberer — only increments when a section is actually emitted,
    sehingga section yang disembunyikan (kosong) tidak meninggalkan nomor bolong."""
    def __init__(self):
        self._n = 0
    def h(self, title: str) -> str:
        self._n += 1
        return f"{self._n}. {title}"


def _is_blank(v: Any) -> bool:
    """True jika konten section dianggap kosong (tidak perlu ditampilkan di report)."""
    if v is None:
        return True
    if isinstance(v, str):
        return v.strip() in ("", "(not provided)", "_not available_", "-")
    if isinstance(v, dict):
        return not any(not _is_blank(x) for x in v.values())
    if isinstance(v, (list, tuple, set)):
        return len(v) == 0
    return False


def _doc_phase_header(doc: Document, phase_name: str, is_draft: bool = False):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = phase_name.upper() + ("  [DRAFT]" if is_draft else "")
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    pPr = p._p.get_or_add_pPr()
    # Background colour
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "1e3a5f")
    shd.set(qn("w:val"),  "clear")
    pPr.append(shd)
    # Outline level 0 = Heading 1 equivalent → collapsible in Word
    outlineLvl = OxmlElement("w:outlineLvl")
    outlineLvl.set(qn("w:val"), "0")
    pPr.append(outlineLvl)


def _doc_body(doc: Document, text: Any):
    t = str(text or "").strip()
    if not t or t in ("(not provided)", "_not available_", "-"):
        p = doc.add_paragraph("(not provided)")
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    else:
        doc.add_paragraph(t)


def _doc_bullet(doc: Document, text: Any):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(str(text))


def _doc_kv(doc: Document, label: str, value: Any):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(str(value or "(not provided)"))


def _doc_styled_table(doc: Document, headers: List[str], rows: List[List],
                      header_color: str = "1e3a5f"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        tc   = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), header_color)
        shd.set(qn("w:val"),  "clear")
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val or "")
            if cells[i].paragraphs[0].runs:
                cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    return t


def _approval_row_data(role: str, action, timestamp) -> List[str]:
    icon  = "✅" if action == "approve" else ("❌" if action == "reject" else "⏳")
    label = "Approved" if action == "approve" else ("Rejected" if action == "reject" else "Pending")
    ts    = f" — {timestamp[:10]}" if timestamp else ""
    return [role, f"{icon} {label}{ts}"]


# ══════════════════════════════════════════════════════════════════════════════
# 3.  SHARED SECTION BUILDERS
# ══════════════════════════════════════════════════════════════════════════════

_GATE_COLOR = {
    "PASS":        ("0f5132", "✅ PASSED"),
    "PASSED":      ("0f5132", "✅ PASSED"),
    "WEAK_PASS":   ("7c4700", "⚠️ CONDITIONAL"),
    "CONDITIONAL": ("7c4700", "⚠️ CONDITIONAL"),
    "FAIL":        ("842029", "❌ FAILED"),
    "FAILED":      ("842029", "❌ FAILED"),
}


def _write_gate_status_line(doc: Document, gate_status: str):
    hex_c, label = _GATE_COLOR.get(gate_status, ("374151", gate_status or "—"))
    p   = doc.add_paragraph()
    run = p.add_run(f"Gate Status: {label}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(int(hex_c[0:2], 16), int(hex_c[2:4], 16), int(hex_c[4:6], 16))


def _effective_gate_status(pid: str, phase: str, raw_status: str) -> str:
    """Return 'PASS' if phase is approved (can_advance=True), else raw_status."""
    try:
        from app.utils import memory as _mem
        appr = _mem.get_phase_approval_status(pid, phase) if pid else {}
        if appr.get("can_advance"):
            return "PASS"
    except Exception:
        pass
    return raw_status


def _write_gate_details(doc: Document, gate: Dict):
    """Write gate policy, missing evidence, risks, next actions."""
    if not gate:
        return
    policy = gate.get("policy") or ""
    if policy:
        p = doc.add_paragraph()
        p.add_run("Policy: ").bold = True
        p.add_run(policy)

    missing = gate.get("missing_evidence") or gate.get("missing_docs") or []
    if isinstance(missing, str):
        missing = [missing]
    if missing:
        doc.add_paragraph().add_run("Missing Evidence:").bold = True
        for m in missing:
            _doc_bullet(doc, m)

    risks = gate.get("risk_if_proceed") or gate.get("risks_if_no_doc") or []
    if isinstance(risks, str):
        risks = [risks]
    if risks:
        doc.add_paragraph().add_run("Risk if Proceed:").bold = True
        for r in risks:
            _doc_bullet(doc, r)

    actions = gate.get("next_actions") or []
    if isinstance(actions, str):
        actions = [actions]
    if actions:
        doc.add_paragraph().add_run("Next Actions:").bold = True
        for a in actions:
            _doc_bullet(doc, a)

    # DEFINE-style: reasons list
    reasons = gate.get("reasons") or []
    if reasons and not (missing or risks or actions):
        doc.add_paragraph().add_run("Gate Findings:").bold = True
        for r in (reasons if isinstance(reasons, list) else [reasons]):
            msg = r.get("message") or r.get("rule") or str(r) if isinstance(r, dict) else str(r)
            _doc_bullet(doc, msg)


def _write_approval_panel(doc: Document, pid: str, phase: str):
    """Write approval status table with names, path-aware."""
    try:
        from app.utils import memory as _mem
        appr = _mem.get_phase_approval_status(pid, phase) if pid else {}
    except Exception:
        appr = {}

    auto_advance = appr.get("auto_advance", False)
    required     = appr.get("required_approvers", [])
    can_advance  = appr.get("can_advance", False)

    # ── Lookup display_name dari secrets.toml berdasarkan role ────────────
    _role_display: Dict[str, str] = {}
    try:
        from app.utils import auth as _auth
        for _udata in _auth._get_users().values():
            _r = _udata.get("role", "")
            if _r and _r not in _role_display:
                _role_display[_r] = _udata.get("display_name", _r)
    except Exception:
        pass

    # ── Project Leader name: project_meta → secrets.toml ─────────────────
    pl_name = ""
    try:
        from app.utils import memory as _mem2
        _meta   = _mem2.load_project_meta(pid) if pid else {}
        pl_name = (_meta.get("project_leader_name", "")
                   or _role_display.get("project_leader", ""))
    except Exception:
        pl_name = _role_display.get("project_leader", "")

    if auto_advance:
        p   = doc.add_paragraph()
        run = p.add_run("Approval: Auto-advance (Quick path — tidak memerlukan approval manual)")
        run.italic = True
        run.font.color.rgb = RGBColor(0x0F, 0x51, 0x32)
        if pl_name:
            _doc_styled_table(doc, ["Role", "Nama", "Keterangan"],
                              [["Project Leader", pl_name, "Penyusun Dokumen"]],
                              header_color="374151")
        return

    # ── Build rows: Project Leader first, then approvers ─────────────────
    rows = []
    if pl_name:
        rows.append(["Project Leader", pl_name, "Penyusun Dokumen", "—"])

    for role_label, key in [("Reviewer", "reviewer"), ("Champion", "champion")]:
        if key in required:
            data         = appr.get(key) or {}
            action       = data.get("action")
            # stored display_name → secrets.toml lookup → role_label fallback
            display_name = (data.get("display_name")
                            or _role_display.get(key, "")
                            or role_label)
            ts           = data.get("timestamp", "")
            icon  = "✅ Approved" if action == "approve" else ("❌ Rejected" if action == "reject" else "⏳ Pending")
            date  = ts[:10] if ts else "—"
            rows.append([role_label, display_name, icon, date])

    if not rows and not required:
        _doc_body(doc, "(Approval tidak diperlukan untuk fase ini)")
        return

    if rows:
        _doc_styled_table(doc, ["Role", "Nama", "Status", "Tanggal"],
                          rows, header_color="374151")

    # Notes
    for role_label, key in [("Reviewer", "reviewer"), ("Champion", "champion")]:
        if key in required:
            data = appr.get(key) or {}
            if data.get("note"):
                p = doc.add_paragraph()
                p.add_run(f"{role_label} Note: ").bold = True
                p.add_run(data["note"])

    p   = doc.add_paragraph()
    run = p.add_run("✅ Gate APPROVED" if can_advance else "⏳ Menunggu approval")
    run.bold = True
    run.font.color.rgb = (RGBColor(0x0F, 0x51, 0x32) if can_advance
                          else RGBColor(0x7C, 0x47, 0x00))


def _write_cover(doc: Document, project_name: str, generated_for_phase: str,
                 is_draft: bool = False):
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DMAIC PROJECT REPORT")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(project_name or "DMAIC Project").font.size = Pt(14)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(
        f"Generated: {datetime.datetime.now().strftime('%d %B %Y')}"
        f"  |  Up to: {generated_for_phase.upper()} Phase"
    ).font.size = Pt(10)

    if is_draft:
        doc.add_paragraph()
        p_d = doc.add_paragraph()
        p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_d = p_d.add_run("⚠️  DRAFT — NOT YET FINALIZED  ⚠️")
        r_d.bold = True
        r_d.font.size = Pt(12)
        r_d.font.color.rgb = RGBColor(0x84, 0x20, 0x29)

    doc.add_paragraph()


def _write_toc(doc: Document, phases: List[str]):
    """Add a Table of Contents page using Word's TOC field (auto-updates on open)."""
    # TOC heading
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("TABLE OF CONTENTS")
    r_title.bold = True
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    doc.add_paragraph()

    # Word TOC field — auto-populates page numbers and headings on open/update
    para = doc.add_paragraph()
    run  = para.add_run()
    # fldChar: begin
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    fldChar_begin.set(qn("w:dirty"), "true")
    run._r.append(fldChar_begin)
    # instrText: TOC instruction (levels 1 only = phase headers)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-1" \\h \\z \\u '
    run._r.append(instrText)
    # fldChar: separate
    fldChar_sep = OxmlElement("w:fldChar")
    fldChar_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar_sep)
    # placeholder lines (replaced when Word updates fields)
    for ph in phases:
        ph_para = doc.add_paragraph()
        ph_para.add_run(f"{ph.upper()} PHASE").bold = True
    # fldChar: end
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar_end)

    doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 4.  DB LOADER
# ══════════════════════════════════════════════════════════════════════════════

def _load_phase_state(pid: str, phase: str) -> Optional[Dict]:
    """Load final, falling back to draft, for the given phase."""
    try:
        from app.utils import memory as _mem
        for suffix in ("final", "draft"):
            loader = getattr(_mem, f"load_{phase}_{suffix}", None)
            if callable(loader):
                state = loader(pid)
                if state:
                    return state
    except Exception:
        pass
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 5.  PHASE SECTION WRITERS
# ══════════════════════════════════════════════════════════════════════════════

def _write_define_section(doc: Document, define_state: Dict, pid: str):
    inputs  = define_state.get("inputs")  or {}
    outputs = define_state.get("outputs") or {}
    gate    = define_state.get("gate_result") or {}
    n       = _normalize_define(define_state)
    _is_draft = define_state.get("status", "") != "final"

    _doc_phase_header(doc, "DEFINE PHASE", is_draft=_is_draft)
    doc.add_paragraph()

    # ── Status & Approval ─────────────────────────────────────────────────
    _doc_heading(doc, "Status & Approval")
    _write_gate_status_line(doc, _effective_gate_status(pid, "define", gate.get("status", "-")))
    _write_approval_panel(doc, pid, "define")
    doc.add_paragraph()


    _sn = _SecNum()

    # ── Business Case ─────────────────────────────────────────────────
    if not _is_blank(n["business_case"]):
        _doc_heading(doc, _sn.h("Business Case"))
        _doc_body(doc, n["business_case"])

    # ── Problem Statement ─────────────────────────────────────────────
    if not _is_blank(n["problem"]):
        _doc_heading(doc, _sn.h("Problem Statement"))
        _doc_body(doc, n["problem"])

    # ── Goal Statement ────────────────────────────────────────────────
    if not _is_blank(n["goal"]):
        _doc_heading(doc, _sn.h("Goal Statement"))
        _doc_body(doc, n["goal"])

    # ── Project Charter ───────────────────────────────────────────────
    charter = n.get("charter") or {}
    if charter and charter.get("available") is not False and any([
        charter.get("project_leader"), charter.get("champion"), charter.get("process_owner"),
    ]):
        _doc_heading(doc, _sn.h("Project Charter"))
        res_rows = []
        for role_label, key in [
            ("Project Leader",   "project_leader"),
            ("Project Champion", "champion"),
            ("Mentoring BB",     "mentor_bb"),
            ("Process Owner",    "process_owner"),
            ("Controller",       "controller"),
        ]:
            val = str(charter.get(key) or "").strip()
            if val and val not in ("TBD", "nan", ""):
                res_rows.append([role_label, val])
        if res_rows:
            _doc_styled_table(doc, ["Role", "Name"], res_rows)

        team = charter.get("team_members") or []
        valid_team = [t for t in team if isinstance(t, dict) and t.get("name")]
        if valid_team:
            doc.add_paragraph().add_run("Team Members:").bold = True
            _doc_styled_table(doc, ["Name", "Function"],
                              [[t.get("name", ""), t.get("function", "")] for t in valid_team])

        timeline = charter.get("timeline") or {}
        if timeline:
            doc.add_paragraph().add_run("Project Timeline:").bold = True
            label_map = {
                "kickoff": "Kick-off", "define_end": "Define",
                "measure_end": "Measure", "analyze_end": "Analyze",
                "improve_end": "Improve", "implement_end": "Implement",
                "control_end": "Control",
            }
            tl_rows = [
                [label_map.get(k, k.replace("_", " ").title()), str(v)]
                for k, v in timeline.items() if v and str(v).strip()
            ]
            if tl_rows:
                _doc_styled_table(doc, ["Milestone", "Target Date"], tl_rows)

    # ── Scope ─────────────────────────────────────────────────────────────
    scope_in  = n["scope_in"]
    scope_out = n["scope_out"]
    _scope_in_list  = scope_in if isinstance(scope_in, list) else ([scope_in] if scope_in else [])
    _scope_out_list = scope_out if isinstance(scope_out, list) else ([scope_out] if scope_out else [])
    if _scope_in_list or _scope_out_list:
        _doc_heading(doc, _sn.h("Scope"))
        p = doc.add_paragraph()
        p.add_run("In Scope:").bold = True
        for item in _scope_in_list:
            _doc_bullet(doc, item)
        p2 = doc.add_paragraph()
        p2.add_run("Out of Scope:").bold = True
        for item in _scope_out_list:
            _doc_bullet(doc, item)

    # ── VOC/VOB → CTQ/CTB ────────────────────────────────────────────────
    voc_table = inputs.get("voc_table") or []
    if voc_table:
        _doc_heading(doc, _sn.h("VOC / VOB → CTQ / CTB"))
        rows = [
            [r.get("Type", ""), r.get("Voice", ""), r.get("Key Issue", ""), r.get("CTQ/CTB", "")]
            for r in voc_table
        ]
        _doc_styled_table(doc, ["Type", "Voice", "Key Issue", "CTQ / CTB"], rows)

    # ── CTQ / CTB List ────────────────────────────────────────────────────
    ctqs = n["ctqs"]
    if isinstance(ctqs, list) and ctqs:
        rows = []
        for c in ctqs:
            if isinstance(c, dict):
                rows.append([c.get("name", ""), c.get("metric", ""),
                             c.get("unit", ""), c.get("description", "")])
        if rows:
            _doc_heading(doc, _sn.h("CTQ / CTB List"))
            _doc_styled_table(doc, ["CTQ Name", "Metric", "Unit", "Description"], rows)

    # ── Y Variable ────────────────────────────────────────────────────────
    if not _is_blank(n["y_variable"]):
        _doc_heading(doc, _sn.h("Y Variable"))
        _doc_body(doc, n["y_variable"])

    # ── SIPOC ─────────────────────────────────────────────────────────────
    sipoc = n["sipoc"] or {}
    if sipoc:
        keys = ["Suppliers", "Inputs", "Process", "Outputs", "Customers"]
        sipoc_lists = []
        for k in keys:
            val = sipoc.get(k) or sipoc.get(k.lower()) or []
            sipoc_lists.append(val if isinstance(val, list) else ([str(val)] if val else []))
        max_len = max([len(x) for x in sipoc_lists], default=0)
        if max_len > 0:
            _doc_heading(doc, _sn.h("SIPOC"))
            rows = [[lst[i] if i < len(lst) else "" for lst in sipoc_lists]
                    for i in range(max_len)]
            _doc_styled_table(doc, keys, rows)

    # ── Benefit Estimate ─────────────────────────────────────────────────
    benefit          = n.get("benefit") or {}
    benefit_from_upload = n.get("benefit_from_upload", False)
    if benefit and isinstance(benefit, dict) and not _is_blank(benefit):
        _doc_heading(doc, _sn.h("Benefit Estimate"))
        if benefit.get("benefit_potentials"):
            _doc_kv(doc, "Benefit Potentials", benefit["benefit_potentials"])
        if benefit.get("calculation_method"):
            _doc_kv(doc, "Calculation Method", benefit["calculation_method"])
        kpi_rows_raw = benefit.get("kpi_table") or []
        if kpi_rows_raw:
            doc.add_paragraph().add_run("KPI Table:").bold = True
            kpi_hdr = ["No", "KPI", "Period", "Unit", "Baseline", "Target", "Δ%",
                       "Agreed by Controlling"]
            kpi_rows = []
            for r in kpi_rows_raw:
                kpi_rows.append([
                    str(r.get("no", "")), r.get("kpi", ""), r.get("period", ""),
                    r.get("unit", ""), r.get("baseline", ""), r.get("target", ""),
                    r.get("delta_pct", ""), r.get("agreed_by_controlling", ""),
                ])
            _doc_styled_table(doc, kpi_hdr, kpi_rows, header_color="0f5132")
        assumptions = benefit.get("assumptions") or []
        if assumptions:
            doc.add_paragraph().add_run("Assumptions:").bold = True
            for a in assumptions:
                _doc_bullet(doc, a)
        # Soft benefits: only from user upload, never from LLM
        if benefit_from_upload:
            soft = benefit.get("soft_benefits") or []
            if soft:
                doc.add_paragraph().add_run("Soft Benefits:").bold = True
                for s in soft:
                    _doc_bullet(doc, s)


def _write_measure_section(doc: Document, measure_state: Dict, pid: str):
    inputs   = measure_state.get("inputs")  or {}
    outputs  = measure_state.get("outputs") or {}
    gate     = measure_state.get("gate")    or {}
    _is_draft = measure_state.get("status", "") != "final"

    _doc_phase_header(doc, "MEASURE PHASE", is_draft=_is_draft)
    doc.add_paragraph()

    # ── Status & Approval ─────────────────────────────────────────────────
    _doc_heading(doc, "Status & Approval")
    _write_gate_status_line(doc, _effective_gate_status(pid, "measure", gate.get("status", "-")))
    _write_approval_panel(doc, pid, "measure")
    doc.add_paragraph()


    _sn = _SecNum()

    # ── Output Measurement ─────────────────────────────────────────────────
    y_conf = outputs.get("y_variable_confirmed") or inputs.get("y_variable") or "-"
    primary = outputs.get("primary_measurement_output") or {}
    if (not _is_blank(y_conf)) or (primary and primary.get("metric_name")):
        _doc_heading(doc, _sn.h("Output Measurement"))
        _doc_kv(doc, "Y Variable Confirmed", y_conf)
        if primary and primary.get("metric_name"):
            _doc_kv(doc, "Primary Measurement", primary.get("metric_name"))
            _doc_kv(doc, "Basis",               primary.get("basis"))
            _doc_kv(doc, "Rationale",           primary.get("rationale"))

    # ── Operational Definition ─────────────────────────────────────────────
    op_defs = outputs.get("operational_definitions") or []
    if op_defs:
        rows = []
        for od in op_defs:
            if isinstance(od, dict):
                rows.append([
                    od.get("metric_name", ""),
                    od.get("what_measured", ""),
                    od.get("method_formula", ""),
                    od.get("unit", ""),
                    od.get("frequency", ""),
                    od.get("decision_criteria", ""),
                ])
        if rows:
            _doc_heading(doc, _sn.h("Operational Definition"))
            _doc_styled_table(doc,
                              ["Metric Name", "What Measured", "Formula/Method",
                               "Unit", "Frequency", "Decision Criteria"],
                              rows)

    # ── Measurement System Analysis (MSA) ─────────────────────────────────
    msa = outputs.get("measurement_system_analysis") or {}
    if msa and isinstance(msa, dict) and not _is_blank(msa):
        _doc_heading(doc, _sn.h("Measurement System Analysis (MSA)"))
        if msa.get("approach"):
            _doc_kv(doc, "Measurement Approach", msa["approach"])
        for field, label in [
            ("error_sources",           "Possible Error Sources"),
            ("cross_checks",            "Cross-checks Performed"),
            ("improvement_steps",       "Improvement Steps"),
            ("alternative_data_sources","Alternative Data Sources"),
        ]:
            items = msa.get(field) or []
            if items:
                doc.add_paragraph().add_run(f"{label}:").bold = True
                for item in items:
                    _doc_bullet(doc, item)
        if msa.get("data_quality_rating"):
            _doc_kv(doc, "Data Quality Rating", msa["data_quality_rating"].upper())
        if msa.get("data_quality_rationale"):
            _doc_kv(doc, "Rationale", msa["data_quality_rationale"])

    # ── Data Collection Plan ───────────────────────────────────────────────
    mplan = outputs.get("measurement_plan") or outputs.get("data_availability_map") or []
    if mplan and isinstance(mplan, list):
        rows = []
        for m in mplan:
            if isinstance(m, dict):
                rows.append([
                    m.get("what", ""), m.get("where", ""), m.get("how", ""),
                    m.get("frequency", ""), m.get("owner", ""),
                ])
        if rows:
            _doc_heading(doc, _sn.h("Data Collection Plan"))
            _doc_styled_table(doc, ["What", "Where", "How", "Frequency", "Owner"], rows)

    # ── Process Performance ────────────────────────────────────────────────
    perf_raw = outputs.get("process_performance_summary")
    perf     = perf_raw if isinstance(perf_raw, dict) else {}
    summary  = perf.get("summary") or {}

    if summary and isinstance(summary, dict) and summary.get("available"):
        _doc_heading(doc, _sn.h("Process Performance"))
        target_info = perf.get("target_info") or {}
        if target_info.get("value") is not None:
            t_val = target_info["value"]
            t_dir = target_info.get("direction", ">=")
            _doc_kv(doc, f"Target ({'min' if t_dir == '>=' else 'max'})", str(t_val))

        stat_rows = []
        for stat_key, label in [
            ("count", "N"), ("mean", "Mean"), ("median", "Median"),
            ("std", "Std Dev"), ("min", "Min"), ("max", "Max"),
            ("p25", "P25"), ("p75", "P75"),
        ]:
            val = summary.get(stat_key)
            if val is not None:
                stat_rows.append([label, str(round(float(val), 4))])
        if stat_rows:
            _doc_styled_table(doc, ["Statistic", "Value"], stat_rows, header_color="0f5132")

        # Chart: jenis dinamis (mesin chart bersama) atau performance bar default
        _mc = perf.get("measure_chart") or {}
        _mc_type = _mc.get("chart_type", "performance")
        chart_b64 = perf.get("chart_b64")
        if _mc_type != "performance" and _mc.get("available"):
            try:
                from app.utils import charts as _charts_m
                import base64, io
                _png_b64 = _charts_m.chart_to_png_b64(_mc, title="Process Performance")
                if _png_b64:
                    doc.add_picture(io.BytesIO(base64.b64decode(_png_b64)), width=Inches(5.5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                _doc_body(doc, "(Chart could not be embedded)")
        elif chart_b64:
            try:
                import base64, io
                chart_bytes  = base64.b64decode(chart_b64)
                chart_stream = io.BytesIO(chart_bytes)
                doc.add_picture(chart_stream, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                _doc_body(doc, "(Chart could not be embedded)")

        commentary = outputs.get("chart_commentary", "")
        if commentary:
            doc.add_paragraph()
            p_c = doc.add_paragraph()
            p_c.add_run("Performance Commentary: ").bold = True
            p_c.add_run(commentary)

    # ── Data Quality Results ───────────────────────────────────────────────
    dq = outputs.get("data_quality_results") or {}
    if dq.get("available"):
        _doc_heading(doc, _sn.h("Data Quality Results"))
        # Completeness
        comp = (dq.get("completeness") or {}).get("columns_with_missing_over_5pct") or []
        if comp:
            _doc_kv(doc, "Completeness", f"Missing >5%: {', '.join(str(c) for c in comp)}")
        else:
            _doc_kv(doc, "Completeness", "✅ Tidak ada kolom dengan missing >5%")

        # Consistency
        cons = (dq.get("consistency") or {}).get("issues") or []
        if cons:
            doc.add_paragraph().add_run("Consistency Issues:").bold = True
            for i in cons:
                if isinstance(i, dict):
                    _doc_bullet(doc, f"{i.get('issue','')} pada '{i.get('field','')}'"
                                     f": {i.get('count','')} records")
                else:
                    _doc_bullet(doc, str(i))
        else:
            _doc_kv(doc, "Consistency", "✅ Tidak ada issue yang terdeteksi")

        # Timeliness
        timeliness = dq.get("timeliness") or {}
        if timeliness.get("date_col") and not timeliness.get("note"):
            _doc_kv(doc, "Timeliness",
                    f"Data dari {timeliness.get('min','?')} sampai {timeliness.get('max','?')}")
            if timeliness.get("freshness_days") is not None:
                _doc_kv(doc, "Freshness", f"{timeliness['freshness_days']:.1f} hari sejak data terakhir")
        else:
            _doc_kv(doc, "Timeliness", "Tidak ada kolom tanggal — timeliness tidak dievaluasi")

        # Dataset Profile
        prof = outputs.get("dataset_profile") or {}
        if prof.get("available"):
            _doc_kv(doc, "Dataset Profile",
                    f"Rows: {prof.get('rows')} | Columns: {prof.get('cols')} "
                    f"| Duplicates: {prof.get('duplicate_rows', 0)} "
                    f"| Missing cells: {prof.get('missing_cells', 0)}")

    # ── Problem Validation ─────────────────────────────────────────────────
    conclusion = outputs.get("measure_conclusion") or {}
    if conclusion:
        _doc_heading(doc, _sn.h("Problem Validation"))
        verdict = conclusion.get("problem_validated", "unknown")
        verdict_map = {
            "true":    "✅ CONFIRMED — Data confirms the problem stated in Define",
            "false":   "❌ NOT CONFIRMED — Data does not support the problem statement",
            "unknown": "⏳ INSUFFICIENT DATA — Cannot validate quantitatively",
        }
        p_v = doc.add_paragraph()
        p_v.add_run(verdict_map.get(verdict, f"Verdict: {verdict}")).bold = True
        _doc_kv(doc, "Confidence", conclusion.get("confidence", "-"))
        basis = conclusion.get("basis") or []
        if basis:
            doc.add_paragraph().add_run("Basis:").bold = True
            for b in basis:
                _doc_bullet(doc, b)


def _write_analyze_section(doc: Document, analyze_state: Dict, pid: str):
    inputs   = analyze_state.get("inputs")  or {}
    outputs  = analyze_state.get("outputs") or {}
    gate     = analyze_state.get("gate")    or {}
    upstream = analyze_state.get("upstream") or {}
    _is_draft = analyze_state.get("status", "") != "final"

    _doc_phase_header(doc, "ANALYZE PHASE", is_draft=_is_draft)
    doc.add_paragraph()

    # ── Status & Approval ─────────────────────────────────────────────────
    _doc_heading(doc, "Status & Approval")
    _write_gate_status_line(doc, _effective_gate_status(pid, "analyze", gate.get("status", "-")))
    _write_approval_panel(doc, pid, "analyze")
    doc.add_paragraph()


    _sn = _SecNum()

    # ── Process Map ────────────────────────────────────────────────────────
    pm_steps = outputs.get("process_map_steps") or []
    pm_rows = []
    for s in pm_steps:
        if isinstance(s, dict) and str(s.get("step_name", "")).strip():
            pm_rows.append([
                str(s.get("step_no", "")), s.get("step_name", ""),
                s.get("responsible", ""), s.get("input", ""),
                s.get("output", ""), s.get("notes", ""),
            ])
    if pm_rows:
        _doc_heading(doc, _sn.h("Process Map"))
        _doc_styled_table(doc,
                          ["Step #", "Step Name", "Responsible",
                           "Input", "Output", "Notes / Failure Mode"],
                          pm_rows)
        pm_check = outputs.get("process_map_check") or {}
        if pm_check.get("feedback"):
            _doc_kv(doc, "Agent Check", pm_check["feedback"])
        if pm_check.get("problem_area_steps"):
            _doc_kv(doc, "Problem Area Steps", ", ".join(pm_check["problem_area_steps"]))

    # ── Fishbone Diagram (6M) ──────────────────────────────────────────────
    fishbone = outputs.get("fishbone_inputs") or {}
    fb_rows = []
    for cat, causes in fishbone.items():
        if isinstance(causes, list):
            for c in causes:
                fb_rows.append([cat, str(c)])
        elif causes:
            fb_rows.append([cat, str(causes)])
    if fb_rows:
        _doc_heading(doc, _sn.h("Fishbone Diagram (6M)"))
        _doc_styled_table(doc, ["Category (6M)", "Potential Cause"], fb_rows)
        fb_check = outputs.get("fishbone_check") or {}
        if fb_check.get("feedback"):
            _doc_kv(doc, "Agent Check", fb_check["feedback"])

    # ── Potential Causes (Xn) ──────────────────────────────────────────────
    pot_causes = outputs.get("potential_causes") or []
    pc_rows = []
    for c in pot_causes:
        if isinstance(c, dict):
            pc_rows.append([c.get("xn", ""), c.get("cause", ""),
                            c.get("source", ""), c.get("direct_effect_on_y", "")])
    if pc_rows:
        _doc_heading(doc, _sn.h("Potential Causes (Xn List)"))
        _doc_styled_table(doc, ["Xn", "Cause", "Source", "Direct Effect on Y"], pc_rows)

    # ── Verification Plan ──────────────────────────────────────────────────
    vplan = outputs.get("verification_plan") or []
    vp_rows = []
    for v in vplan:
        if isinstance(v, dict):
            vp_rows.append([
                v.get("xn", ""), v.get("root_cause", ""),
                v.get("input_measurement_x", ""), v.get("relation_to_y", ""),
                v.get("verification_method", ""), v.get("data_needed", ""),
            ])
    if vp_rows:
        _doc_heading(doc, _sn.h("Verification Plan"))
        _doc_styled_table(doc,
                          ["Xn", "Root Cause", "Input Measurement X",
                           "Relation to Y", "Verification Method", "Data Needed"],
                          vp_rows)

    # ── Verification Results (User-filled field data) ──────────────────────
    vresults = outputs.get("verification_results") or []
    vr_rows = []
    for r in vresults:
        if isinstance(r, dict):
            is_rc_vr = r.get("is_root_cause", False)
            vr_rows.append([
                r.get("xn", ""),
                r.get("root_cause", ""),
                r.get("verification_result", ""),
                str(r.get("impact_pct_of_target", "") or ""),
                r.get("impact_level", ""),
                "✅ Yes" if is_rc_vr else "❌ No",
            ])
    if vr_rows:
        _doc_heading(doc, _sn.h("Verification Results"))
        p_vr_note = doc.add_paragraph()
        p_vr_note.add_run(
            "Diisi oleh Project Leader setelah melakukan verifikasi lapangan/data."
        ).italic = True
        p_vr_note.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        _doc_styled_table(doc,
                          ["Xn", "Root Cause", "Verification Result",
                           "Impact % of Target Gap", "Impact Level", "Root Cause?"],
                          vr_rows, header_color="0f5132")

    # ── Detailed Verification Analysis (Agent-generated) ───────────────────
    analyses = outputs.get("verification_analysis") or []
    if analyses:
        _doc_heading(doc, _sn.h("Detailed Verification Analysis"))
        for a in analyses:
            if not isinstance(a, dict):
                continue
            xn    = a.get("xn", "")
            rc    = a.get("root_cause", "")
            is_rc = a.get("is_root_cause", False)
            icon  = "✅" if is_rc else "❌"

            p_h = doc.add_paragraph()
            p_h.add_run(f"{icon} {xn} — {rc}").bold = True

            hyp = a.get("hypothesis_statement", "")
            if hyp:
                p_hy = doc.add_paragraph()
                p_hy.add_run("Hypothesis: ").bold = True
                p_hy.add_run(hyp)

            desc = a.get("graphical_display_description", "")
            if desc:
                doc.add_paragraph(desc)

            impact = a.get("impact_pct_of_target", "")
            if impact:
                _doc_kv(doc, "Impact on KPI", f"{impact}% of gap to target")

            verdict_text = "ROOT CAUSE CONFIRMED" if is_rc else "Not Confirmed as Root Cause"
            p_v = doc.add_paragraph()
            p_v.add_run(f"Verdict: {verdict_text}").bold = True
            doc.add_paragraph()

    # ── Root Cause Summary Table ───────────────────────────────────────────
    rca_table = outputs.get("root_cause_summary_table") or []
    rca_rows = []
    for r in rca_table:
        if isinstance(r, dict):
            rca_rows.append([
                r.get("xn", ""), r.get("root_cause", ""),
                r.get("is_root_cause", ""), r.get("impact", ""),
                r.get("comment", ""),
            ])
    if rca_rows:
        _doc_heading(doc, _sn.h("Root Cause Summary Table"))
        _doc_styled_table(doc,
                          ["Xn", "Root Cause", "Confirmed?", "Impact", "Comment on KPI"],
                          rca_rows)


def _write_improve_section(doc: Document, improve_state: Dict, pid: str):
    inputs   = improve_state.get("inputs")  or {}
    outputs  = improve_state.get("outputs") or {}
    gate     = improve_state.get("gate")    or {}
    upstream = improve_state.get("upstream") or {}
    _is_draft = improve_state.get("status", "") != "final"

    _doc_phase_header(doc, "IMPROVE PHASE", is_draft=_is_draft)
    doc.add_paragraph()

    # ── Status & Approval ─────────────────────────────────────────────────
    _doc_heading(doc, "Status & Approval")
    _write_gate_status_line(doc, _effective_gate_status(pid, "improve", gate.get("status", "-")))
    _write_approval_panel(doc, pid, "improve")
    doc.add_paragraph()


    _sn = _SecNum()

    # ── Confirmed Root Causes (context) ────────────────────────────────────
    confirmed_rc = upstream.get("confirmed_root_causes") or []
    if confirmed_rc:
        _doc_heading(doc, _sn.h("Confirmed Root Causes (from Analyze)"))
        for r in confirmed_rc:
            if isinstance(r, dict):
                _doc_bullet(doc, f"{r.get('xn', '')} — {r.get('root_cause', '')} "
                                 f"(Impact: {r.get('impact', '')})")
        doc.add_paragraph()

    # ── Must Criteria ──────────────────────────────────────────────────────
    must = outputs.get("must_criteria") or []
    mc_rows = []
    for m in must:
        if isinstance(m, dict):
            mc_rows.append([m.get("criterion", ""), m.get("category", ""),
                            m.get("rationale", "")])
    if mc_rows:
        _doc_heading(doc, _sn.h("Solution Must Criteria"))
        _doc_styled_table(doc, ["Criterion", "Category", "Rationale"], mc_rows)

    # ── Potential Solutions ────────────────────────────────────────────────
    solutions = outputs.get("potential_solutions") or []
    sol_rows = []
    for s in solutions:
        if isinstance(s, dict):
            all_met = "✅" if s.get("all_must_criteria_met") else "❌"
            sol_rows.append([
                s.get("solution_id", ""), s.get("idea", ""),
                s.get("addresses_root_cause", ""), all_met,
                s.get("type", ""), s.get("investment_needed", ""),
                s.get("investment_estimate", ""), s.get("comments", ""),
            ])
    if sol_rows:
        _doc_heading(doc, _sn.h("Potential Solutions"))
        _doc_styled_table(doc,
                          ["ID", "Idea", "Root Cause", "All MC Met",
                           "Type", "Investment", "Est. Cost", "Comments"],
                          sol_rows)

    # ── Selected Solutions ─────────────────────────────────────────────────
    selected_list = outputs.get("selected_solutions") or []
    sel_rows = []
    for i, s in enumerate(selected_list, 1):
        if isinstance(s, dict):
            sel_rows.append([
                str(s.get("no", i)),
                s.get("solution", ""),
                s.get("ctq_ctb_impacted", ""),
                s.get("comments", ""),
            ])
    if sel_rows:
        _doc_heading(doc, _sn.h("Selected Solutions"))
        _doc_styled_table(doc,
                          ["No", "Solution", "CTQ/CTB Impacted", "Comments"],
                          sel_rows, header_color="0f5132")

    # ── Future State / Should-Be Process ───────────────────────────────────
    _future_text = (
        outputs.get("should_be_process")
        or outputs.get("improved_process_map_summary")
        or ""
    )
    if not _is_blank(_future_text):
        _doc_heading(doc, _sn.h("Future State / Should-Be Process"))
        _doc_body(doc, _future_text)

    # ── Pilot Plan ─────────────────────────────────────────────────────────
    pilot = outputs.get("pilot_plan") or {}
    if pilot and isinstance(pilot, dict) and not _is_blank(pilot):
        _doc_heading(doc, _sn.h("Pilot Plan"))
        pilot_rows = [
            ["Scope",            pilot.get("scope", "")],
            ["Duration",         pilot.get("duration", "")],
            ["Success Criteria", pilot.get("success_criteria", "")],
            ["Owner",            pilot.get("owner", "")],
            ["Rollback Plan",    pilot.get("rollback_plan", "")],
        ]
        _doc_styled_table(doc, ["Field", "Detail"], pilot_rows, header_color="374151")
        resources = pilot.get("resources_needed") or []
        if resources:
            doc.add_paragraph().add_run("Resources Needed:").bold = True
            for res in resources:
                _doc_bullet(doc, res)

    # ── Implementation Plan ────────────────────────────────────────────────
    impl = outputs.get("implementation_plan") or []
    impl_rows = []
    for i in impl:
        if isinstance(i, dict):
            impl_rows.append([
                str(i.get("step", "")), i.get("action", ""), i.get("owner", ""),
                i.get("timeline", ""), i.get("deliverable", ""), i.get("dependencies", ""),
            ])
    if impl_rows:
        _doc_heading(doc, _sn.h("Implementation Plan"))
        _doc_styled_table(doc,
                          ["Step", "Action", "Owner", "Timeline", "Deliverable",
                           "Dependencies"],
                          impl_rows)

    # ── Communication Plan ─────────────────────────────────────────────────
    comm = outputs.get("communication_plan") or []
    comm_rows = []
    for c in comm:
        if isinstance(c, dict):
            comm_rows.append([
                c.get("who_informed", ""), c.get("about_what", ""),
                c.get("who_informs", ""), c.get("channel", ""),
                c.get("due_date", ""), c.get("status", "pending"),
            ])
    if comm_rows:
        _doc_heading(doc, _sn.h("Communication Plan"))
        _doc_styled_table(doc,
                          ["Who Informed", "About What", "Who Informs",
                           "Channel", "Due Date", "Status"],
                          comm_rows)

    # ── Implementation Status ──────────────────────────────────────────────
    impl_status = outputs.get("implementation_status") or []
    status_rows = []
    for s in impl_status:
        if isinstance(s, dict):
            status_rows.append([
                s.get("solution", ""),
                s.get("status", "not started"),
                s.get("notes", ""),
            ])
    if status_rows:
        _doc_heading(doc, _sn.h("Implementation Status"))
        _doc_styled_table(doc, ["Solution", "Status", "Notes"], status_rows,
                          header_color="374151")

    # ── Risks & Assumptions ────────────────────────────────────────────────
    risks = outputs.get("risks_and_assumptions") or []
    if risks:
        _doc_heading(doc, _sn.h("Risks & Assumptions"))
        for r in risks:
            _doc_bullet(doc, r)


def _write_control_section(doc: Document, control_state: Dict, pid: str):
    inputs   = control_state.get("inputs")  or {}
    outputs  = control_state.get("outputs") or {}
    gate     = control_state.get("gate")    or {}
    upstream = control_state.get("upstream") or {}
    _is_draft = control_state.get("status", "") != "final"

    _doc_phase_header(doc, "CONTROL PHASE", is_draft=_is_draft)
    doc.add_paragraph()

    # ── Status & Approval ─────────────────────────────────────────────────
    _doc_heading(doc, "Status & Approval")
    _write_gate_status_line(doc, _effective_gate_status(pid, "control", gate.get("status", "-")))
    _write_approval_panel(doc, pid, "control")
    doc.add_paragraph()


    _sn = _SecNum()

    # ── Monitoring & Reaction Plan ─────────────────────────────────────────
    mon_plan = outputs.get("monitoring_reaction_plan") or []
    if mon_plan:
        _doc_heading(doc, _sn.h("Monitoring & Reaction Plan"))
        mon_rows = []
        for m in mon_plan:
            if isinstance(m, dict):
                mon_rows.append([
                    m.get("kpi", ""), m.get("target_value", ""),
                    m.get("measurement_frequency", ""), m.get("who_collects", ""),
                    m.get("who_analyzes", ""), m.get("who_sends_report", ""),
                    m.get("who_takes_action", ""), m.get("deviation_threshold", ""),
                ])
        if mon_rows:
            _doc_styled_table(doc,
                              ["KPI", "Target", "Frekuensi", "Pengumpul Data",
                               "Analis", "Pelapor", "Pelaksana Aksi",
                               "Deviation Threshold"],
                              mon_rows)

        doc.add_paragraph().add_run("Reaction Plan (Penelusuran Root Cause):").bold = True
        for m in mon_plan:
            if not isinstance(m, dict):
                continue
            p_kpi = doc.add_paragraph()
            p_kpi.add_run(
                f"KPI: {m.get('kpi', '')} — Pemicu: {m.get('deviation_threshold', '')}"
            ).bold = True
            react_steps = m.get("reaction_steps") or []
            if react_steps:
                react_rows = []
                for step in react_steps:
                    if isinstance(step, dict):
                        _qref = " ".join(
                            x for x in [step.get("rc_ref", ""), step.get("question", "")] if x
                        )
                        react_rows.append([
                            f"{step.get('step', '')} {('(' + _qref + ')') if _qref else ''}".strip(),
                            step.get("check", ""),
                            step.get("if_yes", ""), step.get("if_no", ""),
                        ])
                if react_rows:
                    _doc_styled_table(doc, ["Langkah", "Pengecekan", "Jika YA", "Jika TIDAK"],
                                      react_rows, header_color="374151")
            esc_actions = m.get("escalation_actions") or []
            if esc_actions:
                p_esc = doc.add_paragraph()
                p_esc.add_run(
                    "Jika solusi gagal / semua RC sudah dicek → Eskalasi:"
                ).bold = True
                for _a in esc_actions:
                    doc.add_paragraph(str(_a), style="List Bullet")
            chart_note = m.get("control_chart_note", "")
            if chart_note:
                p_cn = doc.add_paragraph()
                p_cn.add_run("Catatan: ").bold = True
                p_cn.add_run(chart_note)
            doc.add_paragraph()

    # ── Control Plan per CTQ ───────────────────────────────────────────────
    control_plan = outputs.get("control_plan") or []
    cp_rows = []
    for c in control_plan:
        if isinstance(c, dict):
            cp_rows.append([
                c.get("ctq", ""), c.get("metric", ""), c.get("spec_limit", ""),
                c.get("control_method", ""), c.get("measurement_frequency", ""),
                c.get("sample_size", ""), c.get("owner", ""),
                c.get("recording_system", ""),
            ])
    if cp_rows:
        _doc_heading(doc, _sn.h("Control Plan per CTQ"))
        _doc_styled_table(doc,
                          ["CTQ", "Metric", "Spec Limit", "Control Method",
                           "Frequency", "Sample Size", "Owner", "System"],
                          cp_rows)

    # ── Handover Protocol ──────────────────────────────────────────────────
    handover = outputs.get("handover_protocol") or {}
    if handover and not _is_blank(handover):
        _doc_heading(doc, _sn.h("Handover Protocol"))
        status_rows = [
            ["Should-be process set up",
             "✅" if handover.get("should_be_process_set_up") else "⏳ Pending"],
            ["New KPIs agreed",
             "✅" if handover.get("new_kpis_agreed") else "⏳ Pending"],
            ["Continuous monitoring assured",
             "✅" if handover.get("continuous_monitoring_assured") else "⏳ Pending"],
        ]
        _doc_styled_table(doc, ["Handover Item", "Status"], status_rows, header_color="0f5132")

        # ── Template Q&A (7 pertanyaan) ──
        _kpis_note = (handover.get("kpis_monitoring_note")
                      or handover.get("monitoring_mechanism") or "")
        qa_rows = [
            ["1. Should-be process disiapkan (set up)?",
             "Yes" if handover.get("should_be_process_set_up") else "No"],
            ["2. New process KPIs disepakati & continuous monitoring terjamin?", _kpis_note],
            ["3. Siapa & apa yang harus dilatih?", handover.get("training_note", "")],
            ["4. Open issues", handover.get("open_issues_note", "")],
            ["5. Penanggung jawab open issues", handover.get("open_issues_responsible", "")],
            ["6. Comments / catatan tambahan", handover.get("comments", "")],
            ["7. Responsible (penanggung jawab proses)", handover.get("responsible", "")],
        ]
        qa_rows = [[q, (str(a) if str(a).strip() else "-")] for q, a in qa_rows]
        _doc_styled_table(doc, ["Pertanyaan", "Jawaban"], qa_rows, header_color="374151")

        # ── Confirmation sheet (signed Process Owner) ──
        if handover.get("confirmation_sheet_name"):
            _doc_kv(doc, "Handover Confirmation Sheet (signed Process Owner)",
                    handover["confirmation_sheet_name"])

        # ── Backward-compat: draft lama (training_plan / open_issues list) ──
        training = handover.get("training_plan") or []
        if training:
            doc.add_paragraph().add_run("Training Plan (draft lama):").bold = True
            tr_rows = []
            for t in training:
                if isinstance(t, dict):
                    tr_rows.append([t.get("who", ""), t.get("what", ""),
                                    t.get("by_when", ""), t.get("trainer", ""),
                                    t.get("status", "pending")])
            if tr_rows:
                _doc_styled_table(doc, ["Who", "What", "By When", "Trainer", "Status"],
                                  tr_rows, header_color="374151")

        issues = handover.get("open_issues") or []
        if isinstance(issues, list) and issues:
            doc.add_paragraph().add_run("Open Issues (draft lama):").bold = True
            oi_rows = []
            for i in issues:
                if isinstance(i, dict):
                    oi_rows.append([i.get("issue", ""), i.get("responsible", ""),
                                    i.get("due_date", ""), i.get("status", "open"),
                                    i.get("comments", "")])
            if oi_rows:
                _doc_styled_table(doc, ["Issue", "Responsible", "Due Date", "Status", "Comments"],
                                  oi_rows, header_color="374151")

        handover_complete = bool(handover.get("handover_confirmed")
                                 or handover.get("confirmation_sheet_name"))
        p_hc = doc.add_paragraph()
        r_hc = p_hc.add_run("HANDOVER COMPLETE ✅" if handover_complete else "HANDOVER PENDING ⏳")
        r_hc.bold = True
        r_hc.font.color.rgb = (RGBColor(0x0F, 0x51, 0x32) if handover_complete
                                else RGBColor(0x7C, 0x47, 0x00))

    # ── Benefit Validation ─────────────────────────────────────────────────
    benefit = outputs.get("benefit_validation") or {}
    if benefit and benefit.get("validated_benefit"):
        _doc_heading(doc, _sn.h("Benefit Validation"))
        _doc_kv(doc, "Validated Benefit", benefit.get("validated_benefit"))
        _doc_kv(doc, "Benefit Type",      benefit.get("benefit_type"))
        _doc_kv(doc, "Validated by",      benefit.get("validated_by"))
        _doc_kv(doc, "Validation Date",   benefit.get("validation_date"))
        if benefit.get("approval_sheet_name"):
            _doc_kv(doc, "Benefit Approval Sheet", benefit.get("approval_sheet_name"))
        if benefit.get("tracking_sheet_name"):
            _doc_kv(doc, "Benefit Tracking Sheet", benefit.get("tracking_sheet_name"))

    # ── Control Chart ──────────────────────────────────────────────────────
    cc = outputs.get("control_chart") or {}
    _cc_img = outputs.get("control_chart_image") or {}
    if cc.get("available") or _cc_img.get("b64"):
        _doc_heading(doc, _sn.h("Control Chart"))
    if cc.get("available"):
        from app.utils import charts as _charts_doc
        _ct = cc.get("chart_type", "")
        cc_rows = [["Chart Type", _charts_doc.CHART_LABELS.get(_ct, _ct or "I-MR")]]
        for _lbl, _key in (["Center (X̄)", "center"], ["UCL", "ucl"], ["LCL", "lcl"],
                           ["Std Dev", "std"], ["Mean", "mean"], ["Cp", "cp"],
                           ["Cpk", "cpk"], ["Median", "median"], ["N", "n"]):
            _v = cc.get(_key)
            if _v is not None:
                cc_rows.append([_lbl, str(_v)])
        _doc_styled_table(doc, ["Parameter", "Value"], cc_rows, header_color="374151")

        # Embed chart (sesuai chart_type) via mesin chart bersama
        _png_b64 = _charts_doc.chart_to_png_b64(cc, title="Control Chart")
        if _png_b64:
            try:
                import io as _io_cc, base64 as _b64_cc
                _png = _io_cc.BytesIO(_b64_cc.b64decode(_png_b64))
                doc.add_picture(_png, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                _doc_body(doc, "(Control chart could not be embedded)")

    # Chart improved performance yang diupload (gambar), bila ada
    if _cc_img.get("b64"):
        try:
            import io as _io_ci, base64 as _b64_ci
            doc.add_picture(_io_ci.BytesIO(_b64_ci.b64decode(_cc_img["b64"])), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            _doc_body(doc, "(Uploaded chart could not be embedded)")

    # ── Sustain Audit Protocol ─────────────────────────────────────────────
    sustain = outputs.get("sustain_audit_protocol") or {}
    if sustain and isinstance(sustain, dict) and not _is_blank(sustain):
        _doc_heading(doc, _sn.h("Sustain Audit Protocol"))
        _doc_kv(doc, "Audit Frequency", sustain.get("audit_frequency"))
        _doc_kv(doc, "Auditor",         sustain.get("auditor"))
        _doc_kv(doc, "30-Day Check",    sustain.get("30_day_check"))
        _doc_kv(doc, "60-Day Check",    sustain.get("60_day_check"))
        _doc_kv(doc, "90-Day Check",    sustain.get("90_day_check"))
        checklist = sustain.get("checklist_items") or []
        if checklist:
            doc.add_paragraph().add_run("Audit Checklist:").bold = True
            for item in checklist:
                _doc_bullet(doc, item)
        if sustain.get("escalation_if_slipping"):
            _doc_kv(doc, "Escalation if Slipping", sustain["escalation_if_slipping"])

    # ── Lessons Learned ────────────────────────────────────────────────────
    lessons = outputs.get("lessons_learned") or []
    if lessons:
        _doc_heading(doc, _sn.h("Lessons Learned"))
        for l in lessons:
            _doc_bullet(doc, l)

    # ── Project Closure Checklist ──────────────────────────────────────────
    closure = outputs.get("project_closure_checklist") or []
    cl_rows = []
    for c in closure:
        if isinstance(c, dict):
            cl_rows.append([c.get("item", ""), c.get("status", "pending")])
    if cl_rows:
        _doc_heading(doc, _sn.h("Project Closure Checklist"))
        _doc_styled_table(doc, ["Item", "Status"], cl_rows, header_color="374151")


# ══════════════════════════════════════════════════════════════════════════════
# 6.  EXPORT FUNCTIONS  (cumulative — each includes all previous phases)
# ══════════════════════════════════════════════════════════════════════════════

def export_define_to_word(define_state: Dict[str, Any],
                          path: str = "define_report.docx") -> str:
    pid          = define_state.get("project_id", "")
    n            = _normalize_define(define_state)
    project_name = n["project_name"]
    is_draft     = define_state.get("status", "") != "final"

    doc = Document()
    _doc_setup(doc)
    _write_cover(doc, project_name, "Define", is_draft)

    _write_define_section(doc, define_state, pid)

    doc.save(path)
    return path


def export_measure_to_word(measure_state: Dict[str, Any],
                           path: str = "measure_report.docx") -> str:
    pid      = measure_state.get("project_id", "")
    is_draft = measure_state.get("status", "") != "final"

    # Resolve project name from define
    define_state = _load_phase_state(pid, "define")
    project_name = (
        _normalize_define(define_state)["project_name"]
        if define_state else pid or "DMAIC Project"
    )

    doc = Document()
    _doc_setup(doc)
    _write_cover(doc, project_name, "Measure", is_draft)

    if define_state:
        _write_define_section(doc, define_state, pid)
        doc.add_page_break()

    _write_measure_section(doc, measure_state, pid)

    doc.save(path)
    return path


def export_analyze_to_word(analyze_state: Dict[str, Any],
                           path: str = "analyze_report.docx") -> str:
    pid      = analyze_state.get("project_id", "")
    is_draft = analyze_state.get("status", "") != "final"

    define_state  = _load_phase_state(pid, "define")
    measure_state = _load_phase_state(pid, "measure")
    project_name  = (
        _normalize_define(define_state)["project_name"]
        if define_state else pid or "DMAIC Project"
    )

    doc = Document()
    _doc_setup(doc)
    _write_cover(doc, project_name, "Analyze", is_draft)

    if define_state:
        _write_define_section(doc, define_state, pid)
        doc.add_page_break()

    if measure_state:
        _write_measure_section(doc, measure_state, pid)
        doc.add_page_break()

    _write_analyze_section(doc, analyze_state, pid)

    doc.save(path)
    return path


def export_improve_to_word(improve_state: Dict[str, Any],
                           path: str = "improve_report.docx") -> str:
    pid      = improve_state.get("project_id", "")
    is_draft = improve_state.get("status", "") != "final"

    define_state  = _load_phase_state(pid, "define")
    measure_state = _load_phase_state(pid, "measure")
    analyze_state = _load_phase_state(pid, "analyze")
    project_name  = (
        _normalize_define(define_state)["project_name"]
        if define_state else pid or "DMAIC Project"
    )

    doc = Document()
    _doc_setup(doc)
    _write_cover(doc, project_name, "Improve", is_draft)

    if define_state:
        _write_define_section(doc, define_state, pid)
        doc.add_page_break()

    if measure_state:
        _write_measure_section(doc, measure_state, pid)
        doc.add_page_break()

    if analyze_state:
        _write_analyze_section(doc, analyze_state, pid)
        doc.add_page_break()

    _write_improve_section(doc, improve_state, pid)

    doc.save(path)
    return path


def export_control_to_word(control_state: Dict[str, Any],
                           path: str = "control_report.docx") -> str:
    pid      = control_state.get("project_id", "")
    is_draft = control_state.get("status", "") != "final"

    define_state  = _load_phase_state(pid, "define")
    measure_state = _load_phase_state(pid, "measure")
    analyze_state = _load_phase_state(pid, "analyze")
    improve_state = _load_phase_state(pid, "improve")
    project_name  = (
        _normalize_define(define_state)["project_name"]
        if define_state else pid or "DMAIC Project"
    )

    doc = Document()
    _doc_setup(doc)
    _write_cover(doc, project_name, "Control", is_draft)

    if define_state:
        _write_define_section(doc, define_state, pid)
        doc.add_page_break()

    if measure_state:
        _write_measure_section(doc, measure_state, pid)
        doc.add_page_break()

    if analyze_state:
        _write_analyze_section(doc, analyze_state, pid)
        doc.add_page_break()

    if improve_state:
        _write_improve_section(doc, improve_state, pid)
        doc.add_page_break()

    _write_control_section(doc, control_state, pid)

    doc.save(path)
    return path


# ══════════════════════════════════════════════════════════════════════════════
# 7.  MARKDOWN SUMMARIES  (for UI display — unchanged)
# ══════════════════════════════════════════════════════════════════════════════

def build_define_summary_md(define_state: Dict[str, Any]) -> str:
    n = _normalize_define(define_state)

    def _na(x):
        return x if x not in (None, "", [], {}) else "_not available_"

    md: List[str] = []
    md.append(f"## DEFINE PHASE SUMMARY — {n['project_name']}\n")

    md.append("### Problem Statement")
    md.append(str(_na(n["problem"])) + "\n")
    md.append("### Goal Statement")
    md.append(str(_na(n["goal"])) + "\n")
    md.append("### Business Case")
    md.append(str(_na(n["business_case"])) + "\n")

    md.append("### Scope")
    md.append(f"- In Scope: {_as_text(n['scope_in'])}")
    md.append(f"- Out of Scope: {_as_text(n['scope_out'])}\n")

    md.append("### CTQs")
    ctqs = n["ctqs"]
    if isinstance(ctqs, list) and ctqs:
        for c in ctqs:
            if isinstance(c, dict):
                name   = c.get("name") or c.get("ctq") or "CTQ"
                metric = c.get("metric") or ""
                unit   = c.get("unit") or ""
                desc   = c.get("description") or ""
                target = c.get("target") or ""
                mu  = f" ({metric} {unit})".strip() if metric or unit else ""
                tail = f" — {desc}" if desc else ""
                tgt  = f" | Target: {target}" if target else ""
                md.append(f"- **{name}**{mu}{tail}{tgt}".strip())
            else:
                md.append(f"- {c}")
    else:
        md.append("_not available_")
    md.append("")

    md.append("### Y Variable")
    md.append(str(_na(n["y_variable"])) + "\n")

    md.append("### SIPOC")
    md.append(_sipoc_md(n["sipoc"]))
    md.append("")

    return "\n".join(md).strip()


# Backward-compatible alias
def build_define_summary_markdown(define_state: Dict[str, Any]) -> str:
    return build_define_summary_md(define_state)


def build_measure_summary_md(measure_state: Dict[str, Any]) -> str:
    inputs  = measure_state.get("inputs")  or {}
    outputs = measure_state.get("outputs") or {}
    gate    = measure_state.get("gate")    or {}
    pid     = measure_state.get("project_id", "-")

    y_conf = outputs.get("y_variable_confirmed") or inputs.get("y_variable") or "-"
    perf   = (outputs.get("process_performance_summary") or {}).get("summary") or {}
    mean   = perf.get("mean")
    target = (outputs.get("process_performance_summary") or {}).get("target_info") or {}

    md = [f"## MEASURE SUMMARY — {pid}\n"]
    md.append(f"**Gate Status:** {gate.get('status', '-')}\n")
    md.append(f"**Y Confirmed:** {y_conf}")
    if mean is not None:
        md.append(f"**Baseline Mean:** {mean}")
    if target.get("value") is not None:
        md.append(f"**Target:** {target.get('direction', '')} {target.get('value', '')}")

    op_defs = outputs.get("operational_definitions") or []
    if op_defs:
        md.append(f"\n**Operational Definitions:** {len(op_defs)} metric(s) defined.")

    return "\n".join(md).strip()


# Backward-compatible alias
def build_measure_summary_markdown(measure_state: Dict[str, Any]) -> str:
    return build_measure_summary_md(measure_state)
